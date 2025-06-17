from docx import Document
# import shutil
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import datetime


def make_school(doc, school: str):
    # 二号字体22磅
    p = doc.add_heading(school, 0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in p.runs:
        run = set_font(run, '黑体')
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0, 0, 0)  # 设置字体颜色为黑色
    return p

def make_course(doc, course: str):
    # 三号字体16磅
    p = doc.add_paragraph()
    p.add_run(course).underline = True
    p.add_run("实验报告")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in p.runs:
        run = set_font(run, '黑体')
        run.font.size = Pt(16)
        run.bold = False
        run.font.color.rgb = RGBColor(0, 0, 0)  # 设置字体颜色为黑色
    return p

def make_score(doc, score: str = " "):
    # 五号字体10.5磅
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "成绩"
    for paragraph in table.rows[0].cells[1].paragraphs:
            run = paragraph.add_run(" ")
            run.font.color.rgb = RGBColor(255, 0, 0)
    
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    
    # 启用文本环绕
    # 获取或创建表格属性元素 (w:tblPr)
    tbl_pr = table._tbl.xpath(".//w:tblPr")[0]

    # 创建文字环绕属性元素 (w:tblpPr)
    tblp_pr = OxmlElement("w:tblpPr")
    # 设置文字环绕属性
    tblp_pr.set(qn("w:horzAnchor"), "margin")  # 水平锚点为页边距
    tblp_pr.set(qn("w:vertAnchor"), "text")   # 垂直锚点为文本
    
    # 将文字环绕属性添加到表格属性中
    tbl_pr.append(tblp_pr)


    # 设置列宽
    table.rows[0].cells[0].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
    table.rows[0].cells[1].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        cell.width = Inches(0.6)
        for border in ['top', 'left', 'bottom', 'right']:
            cell._tc.get_or_add_tcPr().append(set_border(border))
        cell.paragraphs[0].paragraph_format.space_after = Pt(10)
        cell.paragraphs[0].paragraph_format.space_before = Pt(10)
        cell.paragraphs[0].paragraph_format.line_spacing = Pt(14)  # 单倍行距
        for run in cell.paragraphs[0].runs:
            run = set_font(run, '黑体')
            run.font.size = Pt(10.5)
            # run.font.color.rgb = RGBColor(0, 0, 0)
    return table

def make_info(doc, info: list, col_widths: list = [1.2, 2.8, 1.2, 1.8]):
    '''
    默认创建四列表格，宽度为 1英寸, 3英寸, 1英寸, 2英寸
    '''
    cols = len(col_widths)
    rows = len(info) // len(col_widths)
    table = doc.add_table(rows, cols)

    # 遍历表格的每一行
    for row in table.rows:
        # 遍历每一行的单元格
        for idx, cell in enumerate(row.cells):
            # 设置单元格的文本
            value = info[row._index * cols + idx]
            if isinstance(value, datetime.date):
                value = value.strftime("%Y-%m-%d")  # 根据需要调整日期格式
            row.cells[idx].text = str(value)
            # 设置中文字体
            for run in cell.paragraphs[0].runs:
                run = set_font(run, '黑体')
                # 四号为14磅
                run.font.size = Pt(14)
            # cell.paragraphs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')

            # 表格行距
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.space_before = Pt(12) # 一行为12磅

            # 设置边框宽度
            cell.width = Inches(col_widths[idx])
            # 获取单元格的边框属性
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            # 创建边框元素
            tcBorders = OxmlElement('w:tcBorders')

            # 如果是偶数列，设置下边框
            if (idx + 1) % 2 == 0:
                bottom = set_border('bottom')
                tcBorders.append(bottom)
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            else:
                # 如果是奇数列，移除所有边框
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = set_border(border_name, 'nil')
                    tcBorders.append(border)
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE

            # 将边框元素添加到单元格属性中
            tcPr.append(tcBorders)

    return table

def make_div(doc, line = False):
    """
    设置分割线。效果等同于在 Word 中输入`---`后快捷插入的分割线。
    """
    p = doc.add_paragraph()
    # 设置字体大小为10.5磅
    for run in p.runs:
        run.font.size = Pt(10.5)
    if line:
        # 获取段落属性
        p_pr = p._element.get_or_add_pPr()
        # 创建边框元素
        p_bdr = OxmlElement('w:pBdr')
        # 创建下边框元素并添加到边框元素
        p_bdr.append(set_border('bottom'))

        # 将边框元素添加到段落属性
        p_pr.append(p_bdr)
    return p

def set_font(run, font):
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    return run

def set_border(border_name, val = "single"):
    border = OxmlElement(f'w:{border_name}')
    border.set(qn('w:val'), val)
    border.set(qn('w:sz'), '4')
    border.set(qn('w:space'), '0')
    border.set(qn('w:color'), '000000')
    return border

def set_color(run, color):
    run.font.color.rgb = color
    return run

# 复制文件
# shutil.copyfile('./md4report/assets/md4report.docx', './md4report/assets/test.docx')

# # 打开复制后的docx文件
# doc = Document('./md4report/tests/teacup.docx')

# starter = doc.paragraphs[0]  # 定位到内容开头

# info = ["专业班级", "计科222", "实验日期", "2025.04.03", "姓名", "钟尹泽", "学号", "202215210229", "实验名称", "实验2.空域图像增强", "指导老师", "华国光"]

# meta = {"school": "广州航海学院", "course": "    数字图像处理及应用实验    ", "score": "", "info": info}

# job = [
#     make_school(doc, meta["school"]),
#     make_course(doc, meta["course"]),
#     make_score(doc),
#     make_div(doc),
#     make_div(doc),
#     make_info(doc, meta["info"]),
#     make_div(doc, True),
# ]

# # 将段落插入到文档开头前，相关用法见“lxml”python库。
# for w in job:
#     starter._element.addprevious(w._element)

# # 保存修改后的文档
# doc.save('./md4report/tests/teacup.docx')


def worker(metadata, file):

    doc = Document(file)
    starter = doc.paragraphs[0]  # 定位到内容开头
    info = [
        "专业班级", metadata["专业班级"],
        "实验日期", metadata["实验日期"],
        "姓名", metadata["姓名"],
        "学号", metadata["学号"],
        "实验名称", metadata["实验名称"],
        "指导老师", metadata["指导老师"],
        ]

    meta = {
        "学校": metadata["学校"],
        "课程": metadata["课程"],
        "info": info
    }

    jobs = [
        make_school(doc, meta["学校"]),
        make_course(doc, meta["课程"]),
        make_score(doc),
        make_div(doc),
        make_div(doc),
        make_info(doc, meta["info"]),
        make_div(doc, True),
    ]

    # 将段落插入到文档开头前，相关用法见“lxml”python库。
    for job in jobs:
        starter._element.addprevious(job._element)

    # 保存修改后的文档
    doc.save(os.path.abspath(file))